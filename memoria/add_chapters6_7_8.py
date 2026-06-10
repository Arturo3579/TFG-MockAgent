import sys
sys.path.append('C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\venv\\Lib\\site-packages')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\memoria-tfg.docx')

def add_heading_custom(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

# 6. IMPLEMENTACIÓN DEL FRONTEND
add_heading_custom('6. Implementación del Frontend', 1)
add_heading_custom('6.1. Arquitectura de la aplicación', 2)
add_para("""El frontend de MockAgent.AI es una Single Page Application (SPA) construida con React 19 y Vite. A diferencia de las aplicaciones multi-página tradicionales, una SPA carga una única página HTML y actualiza el contenido dinámicamente mediante JavaScript, sin recargas completas del navegador.

Herramienta de build: Vite
Vite ha sido seleccionado como herramienta de build en lugar de Create React App (CRA) por varias razones técnicas: velocidad de desarrollo (esbuild en Go para pre-bundling), Hot Module Replacement (HMR) instantáneo, build optimizado con Rollup en producción, y configuración sencilla.

Estructura de archivos:
- src/main.jsx: Punto de entrada. Renderiza la aplicación React en el DOM. Envuelve <App> con GoogleOAuthProvider para la autenticación de Google.
- src/App.jsx: Componente raíz que contiene toda la lógica de la aplicación. La navegación se gestiona mediante un estado central vistaActual (string).
- src/index.css: Estilos globales, variables CSS para theming, y clases de utilidad de Tailwind.
- src/App.css: Estilos específicos de componentes y animaciones personalizadas.
- src/assets/: Imágenes y recursos estáticos.
- public/index.html: Template HTML con meta tags y fuentes de Google.

Gestión de estado:
No se utiliza Redux ni Zustand. El estado se gestiona mediante hooks nativos de React: useState, useEffect, useRef, useCallback. El estado global se almacena en localStorage y sessionStorage para persistencia entre sesiones.

Configuración de Axios:
Se crea una instancia global de Axios con baseURL configurada mediante variable de entorno VITE_API_URL (producción) o fallback a localhost:9090 (desarrollo).""")

add_heading_custom('6.2. Diseño de interfaz y experiencia de usuario', 2)
add_para("""El diseño de MockAgent.AI sigue una filosofía de dark mode premium con acentos dorados, inspirada en las interfaces de herramientas de desarrollo modernas y plataformas SaaS de alta conversión.

Sistema de diseño:
- Paleta de colores: Fondo oscuro profundo (#0a0a1e, #0f0f2e) con degradados sutiles. Acento dorado (#C9A96E) para botones primarios, enlaces y elementos destacados. Texto principal en blanco/gris claro y texto secundario en gris medio.
- Tipografía: Inter (cuerpo de texto) y Space Grotesk (títulos, números). Tamaños de fuente responsive mediante clamp().
- Espaciado: Sistema consistente basado en múltiplos de 4px.
- Radios de borde: 12px para inputs y botones, 16px para cards.
- Sombras: Sombras sutiles con tinte dorado para cards y modales.
- Glassmorphism: Efectos de backdrop-blur en cards y modales para crear profundidad visual.

Theming (claro/oscuro):
- El estado isDarkMode se almacena en localStorage.
- Las variables CSS se definen en :root y se alternan mediante clases .dark en el body.
- Transiciones suaves en todos los cambios de color.

Responsive design:
- Desktop: Layout de dos columnas en dashboard, navbar horizontal.
- Tablet: Cards en grid de 2 columnas.
- Móvil: Navbar hamburger, formularios con inputs táctiles (mínimo 44px), chat widget reposicionado.

Accesibilidad:
- Labels asociados a inputs, contraste WCAG AA, estados focus visibles, iconos SVG escalables.""")

add_heading_custom('6.3. Sistema de navegación y vistas', 2)
add_para("""La navegación es condicional mediante el estado vistaActual (string). Las vistas principales son: landing, login, signup, dashboard, pricing, docs, blog, perfil, y comparativas con competidores.

La función setVistaActual(string) actualiza el estado y React re-renderiza el componente raíz con la nueva vista. Se implementa protección de rutas: usuarios sin token son redirigidos a login, y usuarios logueados que acceden a login son redirigidos al dashboard.

La navegación incluye una navbar con logo, links a landing, pricing, docs, blog, y botones de login/signup o perfil/logout según el estado de autenticación.""")

add_heading_custom('6.4. Landing page y estrategia comercial', 2)
add_para("""La landing page sigue un patrón de diseño SaaS probado para maximizar la conversión de visitantes a usuarios registrados:

1. Hero section: Título con hook de valor, subtítulo descriptivo, botón CTA primario (Empezar Gratis), botón secundario (Ver Demo), y gráfico del producto. Fondo con degradado radial dorado sutil.
2. Barra de confianza: Logos de tecnologías utilizadas (React, Spring Boot, Docker, Stripe) y estadísticas de uptime y latencia.
3. Sección de características: Grid de cards con iconos que destacan funcionalidades clave (Mocking Instantáneo, Seguridad JWT, Rate Limiting, Logs de Auditoría).
4. Sección de flujo de trabajo: Pasos numerados que explican cómo usar la plataforma en menos de 3 minutos.
5. Sección de comparativa: Tabla con checkmarks que posiciona MockAgent.AI frente a competidores.
6. Sección de testimonios: Cards con citas de desarrolladores y equipos de IA.
7. Sección de pricing preview: Cards de planes con precios y características.
8. FAQ: Acordeón con preguntas frecuentes sobre registro, límites, seguridad, cancelación y GDPR.
9. Footer: Links a docs, blog, pricing, términos, privacidad, y copyright.
10. CTA flotante: Botón pegajoso en la parte inferior en móvil.

Animaciones: Framer Motion proporciona animaciones de entrada, transiciones suaves entre secciones, y efectos hover en cards.""")

add_heading_custom('6.5. Dashboard y gestión de endpoints', 2)
add_para("""El dashboard es el núcleo operativo de la aplicación, accesible solo para usuarios autenticados.

Layout: Header con logo, indicador de plan, botón de tema, perfil y logout. Sidebar en móvil (drawer deslizable). Área principal con tabs para crear endpoint, listar endpoints, y ver logs.

Creación de endpoint: Formulario con campos para nombre del path, método HTTP, código de estado, y cuerpo de respuesta JSON. Validación en tiempo real de JSON. Botón de crear con toast de éxito.

Lista de endpoints: Tabla con columnas para path (badge color-coded por método), status, cuerpo truncado, fecha, y acciones (copiar URL, editar, eliminar). Búsqueda en tiempo real por nombre de path. Botón de copiar URL al portapapeles con tooltip de confirmación. Modal de confirmación para eliminar.

Visualización de logs: Tabla con timestamp, método, path, status (color-coded: verde 2xx, amarillo 3xx, rojo 4xx/5xx). Botón para ver cuerpo completo en modal. Filtros por método, status, fecha. Botón para limpiar logs antiguos según retención del plan.

Estados vacíos: Ilustraciones y mensajes amigables cuando no hay endpoints o logs.""")

add_heading_custom('6.6. Sistema de pricing y planes', 2)
add_para("""La página de pricing está diseñada para maximizar la conversión de usuarios gratuitos a pagos.

Diseño: Toggle mensual/anual con animación de Framer Motion. Selector de moneda (EUR, USD, GBP). Grid de 3 cards: Starter (free), Pro, Premium.

Card Starter: Destacada como Gratis para siempre. Características con checkmarks y crossmarks. Botón Empezar Gratis.
Card Pro: Destacada como Recomendado con borde dorado. Precio mensual/anual. Endpoints ilimitados, 5000 peticiones/día, logs 7 días, soporte email. Botón Upgrade a Pro.
Card Premium: Para equipos avanzados. 10000 peticiones/día, logs 14 días, soporte prioritario. Botón Upgrade a Premium.

Comparativa detallada: Tabla de características vs planes con checkmarks verdes o crossmarks grises.

Integración con Stripe: Botones Upgrade llaman al backend que crea una sesión de Stripe Checkout. El frontend redirige al usuario a la URL de checkout. Tras el pago, Stripe redirige al success_url y el frontend refresca el plan del usuario.""")

add_heading_custom('6.7. Chatbot de asistencia inteligente', 2)
add_para("""El chatbot es un sistema de asistencia basado en reglas (rule-based), no un LLM generativo. Está diseñado para responder preguntas frecuentes sobre la plataforma, guiar al usuario en el onboarding y resolver dudas técnicas básicas.

Componente ChatWidget: Widget flotante en la esquina inferior derecha. Al pulsar, se abre un panel de chat con header, área de mensajes, e input de texto. El panel es draggable en desktop y ocupa pantalla completa en móvil.

Motor de respuestas (getAIResponse): Función pura que recibe la pregunta, contexto, e historial. Normaliza la pregunta a minúsculas. Utiliza pattern matching para determinar la intención: login/registro, pricing/planes, mocking/endpoint, soporte/ayuda, seguridad, Stripe/pago, comparativa. Las respuestas son strings con formato Markdown. Incluye memoria de contexto (lastTopic) para preguntas de seguimiento, sugerencias de chips (botones relacionados), y follow-ups para preguntas vagas.

Animaciones: Mensajes del usuario entran desde la derecha, mensajes del bot desde la izquierda con typing indicator antes de mostrar la respuesta.""")

add_heading_custom('6.8. Integración con el backend', 2)
add_para("""La comunicación frontend-backend se realiza exclusivamente mediante peticiones HTTP a la API REST, gestionadas por Axios.

Configuración: baseURL desde VITE_API_URL. En producción apunta a Railway; en desarrollo a localhost:9090.
Autenticación: Header Authorization: Bearer ${token} en todas las peticiones autenticadas. El token se obtiene de getToken() que busca en sessionStorage primero y luego localStorage.
Manejo de errores: try-catch en todas las llamadas. Error de red muestra mensaje de conexión. Error 401 redirige al login. Toast notifications para éxito, error e información.
Carga: Estados isLoading controlan spinners y deshabilitan botones durante peticiones.""")

print("Capítulo 6 completado...")

# 7. SISTEMA DE MONETIZACIÓN CON STRIPE
add_heading_custom('7. Sistema de Monetización con Stripe', 1)
add_heading_custom('7.1. Arquitectura de pagos y Stripe Checkout', 2)
add_para("""El sistema de monetización se ha diseñado siguiendo las mejores prácticas de Stripe para SaaS, priorizando la seguridad, la simplicidad para el usuario y la minimización de la responsabilidad de cumplimiento PCI DSS.

Stripe Checkout es una página de pago alojada por Stripe que se abre en una nueva pestaña. El usuario introduce sus datos de tarjeta directamente en Stripe, nunca en nuestros servidores. Esto elimina la necesidad de manejar datos sensibles de pago en el backend de MockAgent.AI.

El flujo es:
1. El usuario pulsa Upgrade a Pro en el frontend.
2. El frontend envía POST /api/stripe/checkout al backend con el plan deseado y el token JWT.
3. El backend extrae el email del usuario desde el JWT, busca el usuario, y verifica que no tenga ya una suscripción activa.
4. El backend invoca stripe.checkout.sessions.create() con mode subscription, line_items con el price ID del plan, customer_email, success_url, cancel_url, y metadata con userId y plan.
5. Stripe devuelve un objeto Session con una URL.
6. El backend devuelve la URL al frontend.
7. El frontend redirige al usuario a session.url.
8. El usuario completa el pago en Stripe.
9. Stripe redirige al usuario a success_url.
10. El frontend detecta success=true y muestra mensaje de confirmación.""")

add_heading_custom('7.2. Modelos de suscripción y planes de precios', 2)
add_para("""En el dashboard de Stripe se han configurado dos productos con precios recurrentes mensuales:

Producto MockAgent Pro: Precio $4.99 USD/mes. Características: endpoints ilimitados, 5000 peticiones/día, retención de logs 7 días, soporte por email.
Producto MockAgent Premium: Precio $7.99 USD/mes. Características: 10000 peticiones/día, retención de logs 14 días, soporte prioritario (<24h), acceso a beta features.

Plan Starter (Gratuito): No requiere Stripe. Plan por defecto al registrarse. Límites: 5 endpoints, 100 peticiones/día, logs 24h.

Gestión de suscripciones en BD: El campo User.plan es un enum PlanType que se actualiza según el estado de la suscripción en Stripe.""")

add_heading_custom('7.3. Webhooks para gestión de eventos de billing', 2)
add_para("""Los webhooks son la pieza crítica que permite al backend reaccionar a eventos asíncronos de Stripe. El endpoint POST /api/stripe/webhook está configurado como público en SecurityConfig.

Configuración:
- En el dashboard de Stripe, se configura un endpoint webhook con la URL de producción.
- Eventos escuchados: checkout.session.completed, invoice.payment_succeeded, invoice.payment_failed, customer.subscription.deleted, customer.subscription.updated.
- Stripe genera un signing secret almacenado en STRIPE_WEBHOOK_SECRET.

Procesamiento:
1. El backend recibe la petición POST con el payload JSON.
2. Extrae el header Stripe-Signature.
3. Utiliza stripe.webhooks.constructEvent() para validar la firma.
4. Parsea el evento y actúa según el tipo:
   - checkout.session.completed: Extrae email y metadata.plan. Actualiza User.plan a PRO o PREMIUM.
   - invoice.payment_succeeded: Renueva la suscripción.
   - invoice.payment_failed: Muestra banner en dashboard. Plan temporal con gracia.
   - customer.subscription.deleted: Actualiza User.plan a STARTER.
   - customer.subscription.updated: Aplica upgrade inmediato o downgrade al final del ciclo.

Idempotencia: Cada evento tiene un id único. La actualización de un plan ya actualizado no tiene efectos secundarios.""")

add_heading_custom('7.4. Gestión de planes y actualización de cuentas', 2)
add_para("""El sistema permite dos vías para cambiar de plan:
1. Upgrade vía Stripe Checkout: El usuario completa el pago y el webhook actualiza automáticamente su plan.
2. Customer Portal: Stripe proporciona un portal prefabricado donde los usuarios gestionan su método de pago, ven facturas, y cancelan suscripciones. El backend genera un enlace al portal mediante stripe.billingPortal.sessions.create().

Downgrade y cancelación:
- Si el usuario cancela, Stripe envía customer.subscription.deleted.
- El webhook actualiza el plan a STARTER.
- Si el usuario tenía más de 5 endpoints, los existentes siguen funcionando pero no puede crear nuevos.
- Si tenía logs antiguos, la retención se reduce a 24h y se purgan en la siguiente limpieza.

Pruebas gratuitas: Los productos de Stripe pueden configurarse con trial period. En la implementación actual no se configura trial, pero es extensión trivial.""")

add_heading_custom('7.5. Integración frontend-backend de pagos', 2)
add_para("""La integración en el frontend se reduce a dos botones y una redirección:

Botón Upgrade a Pro: Renderizado en pricing y dashboard. Estilo dorado. Al pulsar, establece isLoadingCheckout(true), envía petición al backend, y redirige a la URL de checkout. Si hay error, muestra toast.

Botón Gestionar suscripción: Visible para usuarios Pro/Premium. Abre el Stripe Customer Portal en nueva pestaña.

Indicador de plan: Badge en el header con color según plan (gris Starter, dorado Pro, púrpura Premium). Si es Starter, muestra banner sutil de upgrade.

Verificación post-pago: Al volver de Stripe (success URL), el frontend detecta success=true, llama a GET /api/auth/me para refrescar el plan, y muestra modal de bienvenida al nuevo plan.""")

print("Capítulo 7 completado...")

# 8. RESULTADOS Y ANÁLISIS
add_heading_custom('8. Resultados y Análisis', 1)
add_heading_custom('8.1. Funcionalidades implementadas y demostrables', 2)
add_para("""Al cierre del proyecto, MockAgent.AI es una aplicación funcional y desplegada en producción con las siguientes capacidades verificables:

Autenticación y usuarios: Registro con email/contraseña validadas, login con JWT, login con Google OAuth 2.0, logout, opción Recuérdame.
Gestión de endpoints: Creación con path, método, status, body JSON. Validación JSON. Límite de 5 endpoints para Starter. Endpoints ilimitados para Pro/Premium. Edición, eliminación, búsqueda, copia de URL.
Motor de mocking: Resolución en /mock/{path} con matching exacto. Respuesta con status y body JSON. Instant Mock sin autenticación (expiración 24h).
Sistema de planes: Starter (gratis), Pro ($4.99/mes), Premium ($7.99/mes). Rate limiting por minuto y diario con respuesta 429.
Logs y auditoría: Registro de timestamp, método, path, status, body. Visualización en tabla con filtros. Retención diferenciada (24h, 7d, 14d).
Monetización: Stripe Checkout para suscripciones. Webhooks para actualización automática de planes. Customer Portal.
Interfaz: Landing page, dashboard, pricing, docs, blog, chatbot, tema claro/oscuro, responsive, toast notifications, modales de confirmación.
Despliegue: Backend en Railway, frontend en Vercel, dominio mockagentai.com, SSL automático, health checks.""")

add_heading_custom('8.2. Métricas de despliegue y rendimiento', 2)
add_para("""Disponibilidad: Uptime del backend >99.5% (monitoreado por Railway). Uptime del frontend >99.9% (Vercel edge network).

Rendimiento:
- Tiempo de respuesta backend (CRUD): ~50-150ms medio.
- Tiempo de respuesta motor de mocking: ~20-80ms medio.
- Build frontend: ~3s (dev), ~15s (producción).
- Build backend: ~45s (sin tests), ~90s (con tests).
- Despliegue Railway: ~2-3 minutos desde push.
- Despliegue Vercel: ~30-60 segundos desde push.

Recursos:
- Backend: Contenedor Docker con 512MB-1GB RAM (Railway free tier).
- Frontend: Static files en CDN (sin coste de servidor).
- Base de datos: TiDB Cloud serverless (5GB free tier).

Tráfico: El proyecto está en fase de validación inicial. No se han realizado campañas de marketing masivo.""")

add_heading_custom('8.3. Comparativa con soluciones existentes', 2)
add_para("""Tabla comparativa:

Característica | MockAgent.AI | Beeceptor | Postman Mock | WireMock | Mockbin
Interfaz web amigable | Sí | Sí | Sí | No (JSON) | Sí (básica)
Autenticación de usuarios | JWT + OAuth | No | Postman account | No | No
OAuth 2.0 Google | Sí | No | No | No | No
Planes de uso (Freemium) | Starter/Pro/Premium | Muy limitado | Team/Enterprise | N/A | No
Rate limiting | Por usuario | No | No | No | No
Logs de auditoría | Con retención | Básicos | Sí | Sí | Sí
Instant Mock (sin registro) | Sí | No | No | No | Sí (bin)
Pasarela de pagos | Stripe | No | No | No | No
Chatbot de asistencia | Reglas | No | No | No | No
Dominio custom | mockagentai.com | No | No | No | No
Orientado a agentes IA | Sí | No | No | No | No
Despliegue real | Railway+Vercel | SaaS | SaaS | Self-hosted | SaaS

MockAgent.AI se diferencia por ser la única solución que combina interfaz web moderna, autenticación dual, sistema de planes freemium con rate limiting, monetización Stripe, logs de auditoría, chatbot contextual, y despliegue real con dominio propio, orientado al nicho de agentes de IA.""")

add_heading_custom('8.4. Casos de uso y validación práctica', 2)
add_para("""Caso de uso 1: Desarrollador de agentes de IA (Startup)
María desarrolla un agente que consulta precios de criptomonedas. La API externa tiene rate limit de 100/día y cuesta $0.01/petición. María crea un mock en MockAgent.AI. Durante 2 semanas realiza 500 peticiones/día de prueba. Coste sin MockAgent.AI: $70 + bloqueos. Con MockAgent.AI: $0 (plan Starter).

Caso de uso 2: Testing de escenarios de error
Carlos necesita probar cómo su agente reacciona ante un error 500. No puede forzarlo en producción. Configura un mock con status 500 y body JSON de error. Ejecuta su agente, verifica el manejo de errores, y despliega con confianza.

Caso de uso 3: Demo comercial sin dependencias
Una startup de IA hace una demo a un inversor. No dependen de APIs de terceros. Configuran un mock con datos realistas. La demo funciona perfectamente sin conexión a internet externa, transmitiendo profesionalidad.

Validación práctica:
Los tests de integración (JUnit 5) validan casos críticos:
- testSignupNewUser: Registro y generación de JWT.
- testLoginSuccess: Flujo de login.
- testCreateEndpoint: Creación autenticada.
- testEndpointStarterLimit: Límite de 5 endpoints (HTTP 403).
- testMockResponse: Motor de mocking responde con JSON configurado (HTTP 418 teapot).""")

print("Capítulo 8 completado...")

output_path = 'C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\memoria-tfg.docx'
doc.save(output_path)
print(f"Documento actualizado en: {output_path}")
