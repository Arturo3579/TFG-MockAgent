import sys
sys.path.append('C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\venv\\Lib\\site-packages')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\memoria-tfg.docx')

def add_heading_custom(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

# 9. DISCUSIÓN
add_heading_custom('9. Discusión', 1)
add_heading_custom('9.1. Fortalezas del proyecto', 2)
add_para("""1. Arquitectura limpia y profesional: La separación de capas (controller-service-repository) facilita el mantenimiento, el testing y la escalabilidad. El uso de DTOs desacopla la API de la base de datos.

2. Stack tecnológico moderno y probado: Spring Boot 3.5 + React 19 + Vite es una combinación actualizada y alineada con las demandas del mercado laboral. TailwindCSS y Framer Motion proporcionan una UI moderna sin CSS excesivo.

3. Autenticación dual robusta: La combinación de JWT + BCrypt + Google OAuth 2.0 con validación server-side proporciona un sistema seguro y flexible. No se almacenan contraseñas en texto plano.

4. Despliegue real y dominio propio: A diferencia de muchos TFGs que se quedan en localhost, MockAgent.AI está desplegado en producción con dominio custom, SSL, y health checks. Esto demuestra competencias reales de DevOps.

5. Modelo de negocio freemium integrado: La implementación de Stripe Checkout, webhooks, y planes diferenciados convierte el proyecto en un producto SaaS viable, no solo en una herramienta de práctica.

6. Chatbot contextual: Aunque es rule-based, el chatbot proporciona valor real al usuario, respondiendo preguntas frecuentes y guiando el onboarding. La memoria de contexto y las sugerencias mejoran la UX.

7. Responsive y accesible: El frontend funciona correctamente en móvil, tablet y desktop. El tema claro/oscuro y el contraste de colores cumplen estándares básicos de accesibilidad.

8. Documentación completa: README, DEPLOY.md, y docs internos en la aplicación proporcionan guías claras para usuarios y desarrolladores.""")

add_heading_custom('9.2. Limitaciones y amenazas identificadas', 2)
add_para("""1. Rate limiting en memoria: RateLimitService utiliza ConcurrentHashMap en memoria. Funciona para un único servidor pero no escala a múltiples instancias. La solución sería migrar a Redis.

2. Sin tests de frontend: No se han implementado tests unitarios ni de integración en el frontend (no Vitest, no Jest, no Cypress). Esto aumenta el riesgo de regressions visuales y funcionales.

3. Sin migraciones de base de datos: El uso de ddl-auto=update es conveniente para desarrollo pero peligroso en producción a largo plazo. No hay control de versiones del esquema (Flyway/Liquibase). Un cambio de modelo podría causar pérdida de datos.

4. Sin RBAC: Todos los usuarios autenticados tienen la misma authority en Spring Security (lista vacía). El control de acceso se basa únicamente en el claim plan del JWT. No hay roles de administrador, moderador, etc.

5. Chatbot rule-based limitado: El chatbot no utiliza un LLM ni procesamiento de lenguaje natural avanzado. Las respuestas son predefinidas y el pattern matching es básico. Puede fallar con preguntas complejas o ambiguas.

6. Sin multi-tenancy ni workspaces: Actualmente, cada usuario gestiona sus endpoints individuales. No hay equipos, workspaces compartidos, ni roles de colaborador. Esto limita la adopción por empresas.

7. Dependencia de TiDB Cloud: La base de datos está en TiDB Cloud (free tier). Si se exceden los límites del plan gratuito, el servicio podría degradarse. No hay un plan de contingencia documentado.

8. SSL gestionado por plataforma: Aunque Railway y Vercel proporcionan SSL automático, no hay control directo sobre los certificados ni configuración de HSTS en el código.

9. Sin sistema de notificaciones: No hay notificaciones por email, SMS, ni push para alertas de límite de plan, expiración de suscripción, o actividad sospechosa.

10. Código frontend monolítico: Todo el frontend está en un único archivo App.jsx de ~2.700 líneas. Aunque funcional, dificulta el mantenimiento a largo plazo. La refactorización a componentes separados sería necesaria.""")

add_heading_custom('9.3. Comparación con el estado del arte', 2)
add_para("""MockAgent.AI se posiciona como una solución de nicho entre las herramientas de mocking existentes. No compite directamente con Postman (suite completa de API development) ni con WireMock (biblioteca de testing para Java). Compite en el espacio de SaaS de mocking ligero (Beeceptor, Mockbin) pero con diferenciación clara: está diseñado específicamente para el desarrollo de agentes de IA y incluye monetización nativa.

Las herramientas enterprise (Postman, WireMock) ofrecen más funcionalidades (contract testing, CI/CD integration, colecciones compartidas) pero requieren más configuración técnica y no incluyen pasarelas de pago ni chatbots. MockAgent.AI sacrifica algunas funcionalidades avanzadas a cambio de simplicidad, onboarding rápido, y modelo de negocio integrado.""")

add_heading_custom('9.4. Escalabilidad y mantenimiento a largo plazo', 2)
add_para("""Escalabilidad técnica: El backend es stateless (no guarda sesiones en memoria), por lo que puede escalar horizontalmente mediante múltiples contenedores Docker detrás de un load balancer. El único cuello de botella es el rate limiting en memoria, que requeriría Redis. El frontend es estático y escala infinitamente mediante CDN. La base de datos TiDB Cloud es serverless y auto-escalable.

Escalabilidad de negocio: El modelo freemium permite adquisición orgánica de usuarios gratuitos que pueden convertirse a pagos. La integración de Stripe permite añadir nuevos planes sin modificar el código (solo configurar productos en Stripe y actualizar el enum PlanType). El dominio propio permite construir marca y SEO.

Mantenimiento: Spring Boot proporciona soporte LTS y actualizaciones de seguridad. React 19 es la última versión estable, pero el ecosistema evoluciona rápidamente. Será necesario mantener las dependencias actualizadas. El monolito actual puede evolucionar a microservicios si el tráfico lo justifica.""")

print("Capítulo 9 completado...")

# 10. CONCLUSIONES
add_heading_custom('10. Conclusiones', 1)
add_heading_custom('10.1. Cumplimiento de objetivos', 2)
add_para("""El proyecto MockAgent.AI ha cumplido satisfactoriamente todos los objetivos específicos planteados:

1. Se ha diseñado e implementado una arquitectura backend por capas con Spring Boot 3.5, Java 17, y Maven, siguiendo buenas prácticas de separación de responsabilidades.

2. Se ha implementado un sistema de autenticación dual (JWT con BCrypt + Google OAuth 2.0 con validación server-side) que garantiza la seguridad y la flexibilidad.

3. Se ha desarrollado un motor de mocking funcional que permite crear, editar, eliminar y consumir endpoints REST con respuestas JSON configurables, incluyendo Instant Mock sin autenticación.

4. Se ha implementado un sistema de planes de uso (Starter, Pro, Premium) con rate limiting por minuto y por día, controlando el acceso a recursos según el nivel de suscripción.

5. Se ha integrado Stripe como pasarela de pagos, implementando Stripe Checkout para suscripciones recurrentes, webhooks para gestión de eventos de billing, y actualización automática de planes.

6. Se ha desarrollado un frontend moderno con React 19, Vite, TailwindCSS, y Framer Motion, incluyendo landing page, dashboard, pricing, docs, blog, chatbot, y tema claro/oscuro.

7. Se ha implementado un sistema de logging y auditoría que registra cada petición con timestamp, método, path, status, y body, aplicando retención diferenciada por plan.

8. Se ha desplegado la aplicación en infraestructura cloud real (Railway para backend, Vercel para frontend, TiDB Cloud para base de datos) con dominio personalizado (mockagentai.com) y certificados SSL.

9. Se ha diseñado y ejecutado un conjunto de 9 tests de integración con JUnit 5 y H2 in-memory que validan los flujos críticos de autenticación, CRUD de endpoints, límite de planes, y respuestas de mocks.""")

add_heading_custom('10.2. Impacto del proyecto y contribución', 2)
add_para("""Desde una perspectiva académica, MockAgent.AI demuestra la capacidad de integrar conocimientos transversales del Grado en Ingeniería Informática: diseño de arquitecturas software, programación orientada a objetos, bases de datos relacionales, desarrollo web frontend, seguridad informática, despliegue de aplicaciones y gestión de proyectos. El hecho de que el proyecto esté desplegado en producción con dominio propio y pasarela de pagos real (Stripe) eleva el trabajo de un ejercicio académico a un producto viable mínimo (MVP).

Desde una perspectiva profesional, el proyecto aporta valor al ecosistema de desarrollo de agentes de IA, proporcionando una herramienta especializada que reduce costes, elimina dependencias de terceros durante el desarrollo, y permite simular escenarios de error de forma controlada. El modelo de negocio freemium con planes diferenciados demuestra una comprensión práctica de la monetización de productos SaaS.

La contribución técnica más relevante es la integración de Google OAuth 2.0 con validación server-side en una arquitectura JWT stateless, manteniendo la coherencia del sistema de autenticación independientemente del método de login. Además, el chatbot de asistencia basado en reglas con memoria de contexto representa una solución innovadora para el onboarding de usuarios en herramientas de desarrollo.""")

add_heading_custom('10.3. Líneas futuras de trabajo', 2)
add_para("""El proyecto abre múltiples líneas de trabajo para futuras iteraciones:

1. Integración de LLM en el chatbot: Migrar el chatbot rule-based a un modelo de lenguaje (OpenAI GPT, Claude, o Llama) mediante API, permitiendo respuestas más naturales y contextuales.

2. Workspaces y colaboración en equipo: Implementar multi-tenancy con workspaces donde múltiples usuarios compartan endpoints, roles (admin, editor, viewer), y actividad en tiempo real.

3. WebSockets y tiempo real: Añadir WebSockets para notificaciones push en el dashboard (nueva petición recibida, límite de plan cercano), y para colaboración en tiempo real en workspaces.

4. Migración de rate limiting a Redis: Reemplazar el ConcurrentHashMap en memoria por Redis para permitir escalado horizontal del backend sin perder consistencia.

5. Migraciones de base de datos: Implementar Flyway o Liquibase para control de versiones del esquema, facilitando despliegues seguros en producción.

6. Tests de frontend: Añadir Vitest para tests unitarios de componentes React, y Cypress o Playwright para tests end-to-end (E2E) de flujos críticos.

7. Análisis avanzado y dashboards: Implementar dashboards analíticos con gráficos de uso (peticiones por día, endpoints más usados, errores comunes) utilizando Chart.js o Recharts.

8. API pública y SDKs: Exponer una API REST pública documentada con OpenAPI/Swagger, y desarrollar SDKs para JavaScript/TypeScript, Python, y Java.

9. Webhooks personalizados: Permitir a los usuarios configurar webhooks propios que se disparen cuando sus endpoints reciben peticiones, integrando con Slack, Discord, o Zapier.

10. Internacionalización (i18n): Añadir soporte para múltiples idiomas (inglés, español, francés, alemán) mediante react-i18next.

11. Password reset y email verification: Implementar flujo de recuperación de contraseña mediante tokens de un solo uso enviados por email, y verificación de email tras registro.

12. Refactorización del frontend: Dividir App.jsx en componentes separados, custom hooks, y un sistema de routing con react-router-dom para permitir compartir URLs de vistas específicas.""")

print("Capítulo 10 completado...")

# REFERENCIAS
add_heading_custom('Referencias Bibliográficas', 1)
refs = [
    "Fielding, R. (2000). Architectural Styles and the Design of Network-based Software Architectures [Tesis doctoral]. University of California, Irvine.",
    "Google. (2024). Google Identity Services: Sign In with Google for Web. Google Developers.",
    "Hibernate. (2024). Hibernate ORM 6.5 Documentation. https://hibernate.org/documentation/",
    "IETF. (2015). RFC 7519: JSON Web Token (JWT). https://datatracker.ietf.org/doc/html/rfc7519",
    "Meta Open Source. (2024). React 19. React Documentation. https://react.dev/",
    "MySQL. (2024). MySQL 8.0 Reference Manual. https://dev.mysql.com/doc/refman/8.0/en/",
    "OWASP. (2023). OWASP Top 10:2021. Open Web Application Security Project. https://owasp.org/Top10/",
    "Spring.io. (2024). Spring Boot 3.5 Reference Documentation. https://docs.spring.io/spring-boot/docs/3.5.14/reference/html/",
    "Spring.io. (2024). Spring Security Reference. https://docs.spring.io/spring-security/reference/",
    "Stripe. (2024). Stripe API Reference. https://stripe.com/docs/api",
    "Stripe. (2024). Stripe Checkout: Quickstart. https://stripe.com/docs/checkout/quickstart",
    "Tailwind CSS. (2024). Tailwind CSS Documentation. https://tailwindcss.com/docs",
    "Vercel. (2024). Vercel Documentation: Deploying Frontend Applications. https://vercel.com/docs",
    "Vite. (2024). Vite Documentation. https://vitejs.dev/guide/"
]
for r in refs:
    add_para(r)

print("Referencias completadas...")

# ANEXOS
add_heading_custom('Anexos', 1)

add_heading_custom('Anexo A. Configuración de seguridad del backend', 2)
add_para("""El archivo SecurityConfig.java centraliza toda la configuración de seguridad de Spring Boot:

@Configuration
@EnableWebSecurity
@RequiredArgsConstructor
public class SecurityConfig {
    private final JwtAuthFilter jwtAuthFilter;

    @Bean
    public SecurityFilterChain securityFilterChain(HttpSecurity http) throws Exception {
        http
            .csrf(AbstractHttpConfigurer::disable)
            .sessionManagement(session -> 
                session.sessionCreationPolicy(SessionCreationPolicy.STATELESS))
            .cors(cors -> cors.configurationSource(corsConfigurationSource()))
            .authorizeHttpRequests(auth -> auth
                .requestMatchers("/", "/api/health", "/api/auth/**", 
                    "/api/instant/**", "/api/stripe/webhook", "/mock/**", "OPTIONS /**")
                .permitAll()
                .anyRequest().authenticated()
            )
            .addFilterBefore(jwtAuthFilter, UsernamePasswordAuthenticationFilter.class);
        return http.build();
    }

    @Bean
    public CorsConfigurationSource corsConfigurationSource() {
        CorsConfiguration config = new CorsConfiguration();
        config.setAllowedOrigins(Arrays.asList(
            "http://localhost:5173", "http://localhost:3000",
            "https://tfg-mockagent.vercel.app",
            "https://mockagent-ai.vercel.app",
            "https://mockagentai.com", "https://www.mockagentai.com"
        ));
        config.setAllowedMethods(Arrays.asList("GET", "POST", "PUT", "DELETE", "OPTIONS"));
        config.setAllowedHeaders(Arrays.asList("Authorization", "Content-Type", "X-Requested-With"));
        config.setAllowCredentials(true);
        config.setMaxAge(3600L);
        UrlBasedCorsConfigurationSource source = new UrlBasedCorsConfigurationSource();
        source.registerCorsConfiguration("/**", config);
        return source;
    }

    @Bean
    public PasswordEncoder passwordEncoder() {
        return new BCryptPasswordEncoder();
    }
}""")

add_heading_custom('Anexo B. Configuración de despliegue en producción', 2)
add_para("""Backend (Railway):
- Dockerfile multi-stage:
  Stage 1 (build): eclipse-temurin:17-jdk-alpine, copia pom.xml y src, ejecuta mvn clean package -DskipTests.
  Stage 2 (runtime): eclipse-temurin:17-jre-alpine, copia el JAR desde stage 1, expone puerto 8080, ejecuta java -jar app.jar.
- Variables de entorno en Railway:
  DATABASE_URL: URL JDBC de TiDB Cloud.
  JWT_SECRET: Clave secreta para firmar JWT (mínimo 256 bits).
  GOOGLE_CLIENT_ID: Client ID de Google OAuth.
  STRIPE_SECRET_KEY: Clave secreta de Stripe (modo live).
  STRIPE_WEBHOOK_SECRET: Signing secret del webhook de Stripe.
  STRIPE_PRICE_PRO_MONTHLY: ID del precio de Stripe para plan Pro.
  STRIPE_PRICE_PREMIUM_MONTHLY: ID del precio de Stripe para plan Premium.
  PORT: Puerto asignado por Railway (8080 por defecto).

Frontend (Vercel):
- Build command: npm run build (vite build).
- Output directory: dist.
- Variables de entorno:
  VITE_API_URL: URL del backend en Railway (https://tfg-mockagent-production.up.railway.app).
  VITE_GOOGLE_CLIENT_ID: Client ID de Google OAuth.
- vercel.json: SPA rewrite rules (todas las rutas redirigen a index.html).

Dominio (Namecheap):
- CNAME record: www -> cname.vercel-dns.com.
- A record: @ -> 76.76.21.21 (IP de Vercel).
- Verificado en dashboard de Vercel.""")

add_heading_custom('Anexo C. Esquema relacional de la base de datos', 2)
add_para("""Tabla users:
- id (BIGINT, PK, AUTO_INCREMENT)
- email (VARCHAR(255), UNIQUE, NOT NULL)
- password (VARCHAR(255), NOT NULL)
- google_id (VARCHAR(255), UNIQUE, NULL)
- plan (ENUM('STARTER','PRO','PREMIUM','ENTERPRISE'), NOT NULL, DEFAULT 'STARTER')

Tabla mock_endpoints:
- id (BIGINT, PK, AUTO_INCREMENT)
- user_id (BIGINT, FK users.id, NULL)
- path (VARCHAR(255), NOT NULL)
- method (VARCHAR(10), NOT NULL)
- status (INT, DEFAULT 200)
- response_body (TEXT, NULL)
- is_demo (BOOLEAN, DEFAULT FALSE)
- created_at (TIMESTAMP, DEFAULT CURRENT_TIMESTAMP)

Tabla request_logs:
- id (BIGINT, PK, AUTO_INCREMENT)
- endpoint_id (BIGINT, FK mock_endpoints.id, NULL)
- user_id (BIGINT, FK users.id, NULL)
- timestamp (TIMESTAMP, DEFAULT CURRENT_TIMESTAMP)
- method (VARCHAR(10), NOT NULL)
- path (VARCHAR(255), NOT NULL)
- status_code (INT, NOT NULL)
- request_body (TEXT, NULL)

Relaciones:
- users 1:N mock_endpoints (un usuario tiene muchos endpoints)
- users 1:N request_logs (un usuario genera muchos logs)
- mock_endpoints 1:N request_logs (un endpoint recibe muchos logs)

Índices:
- users.email (UNIQUE)
- users.google_id (UNIQUE)
- mock_endpoints.user_id (INDEX)
- request_logs.endpoint_id (INDEX)
- request_logs.user_id (INDEX)
- request_logs.timestamp (INDEX)""")

print("Anexos completados...")

output_path = 'C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\memoria-tfg.docx'
doc.save(output_path)
print(f"DOCUMENTO FINALIZADO Y GUARDADO EN: {output_path}")
print("¡La memoria está completa!")
