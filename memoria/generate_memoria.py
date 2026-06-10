import sys
sys.path.append('C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\venv\\Lib\\site-packages')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Estilos
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Función para añadir título
def add_heading_custom(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

# Función para añadir párrafo
def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

# PORTADA
portada = doc.add_paragraph()
portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = portada.add_run('MockAgent.AI\n')
run.bold = True
run.font.size = Pt(24)
run = portada.add_run('Plataforma SaaS de API Mocking para el Desarrollo de Agentes de Inteligencia Artificial\n\n')
run.font.size = Pt(16)
run = portada.add_run('Trabajo de Fin de Grado\n\n')
run.font.size = Pt(14)
run = portada.add_run('Autor: [Nombre del estudiante]\n')
run = portada.add_run('Tutor: [Nombre del tutor]\n')
run = portada.add_run('Grado: Grado en Ingeniería Informática / Ingeniería del Software\n')
run = portada.add_run('Departamento: [Departamento]\n')
run = portada.add_run('Universidad: [Nombre de la universidad]\n')
run = portada.add_run('Fecha: Junio 2026\n')
doc.add_page_break()

# RESUMEN
add_heading_custom('Resumen', 1)
resumen = """El presente trabajo de fin de grado describe el diseño, desarrollo y despliegue de MockAgent.AI, una plataforma SaaS (Software as a Service) especializada en la creación de APIs simuladas (mock APIs) orientada al desarrollo y testing de agentes de inteligencia artificial. En el contexto actual, donde el desarrollo de aplicaciones basadas en IA generativa experimenta un crecimiento exponencial, los desarrolladores enfrentan desafíos significativos relacionados con los costes de consumo de APIs externas, los límites de petición (rate limits), la latencia de red y la imposibilidad de simular escenarios de error de forma controlada.

La plataforma se ha desarrollado siguiendo una arquitectura moderna de aplicación web full-stack. El backend se implementa utilizando el framework Spring Boot 3.5 con Java 17, adoptando un patrón de arquitectura por capas (controller-service-repository) que garantiza la separación de responsabilidades y la mantenibilidad del código. El sistema de autenticación implementa un doble mecanismo: autenticación clásica basada en tokens JWT (JSON Web Tokens) con hash de contraseñas mediante BCrypt, e integración de inicio de sesión social mediante OAuth 2.0 de Google, validando los tokens de identidad en el servidor mediante la librería oficial de Google. La persistencia de datos se gestiona a través de Spring Data JPA sobre una base de datos MySQL 8.0, desplegada en la nube mediante TiDB Cloud.

El frontend se ha construido como una Single Page Application (SPA) utilizando React 19 con Vite como herramienta de build, lo que proporciona tiempos de compilación ultrarrápidos y una experiencia de desarrollo optimizada. El diseño de interfaz emplea TailwindCSS para estilos utility-first, garantizando un diseño totalmente responsive y adaptativo a dispositivos móviles.

En cuanto a la monetización, la plataforma integra Stripe como pasarela de pagos, permitiendo la suscripción a planes Pro y Premium mediante Stripe Checkout. El sistema de webhooks de Stripe gestiona los eventos de suscripción, cancelación y renovación, actualizando automáticamente el plan del usuario en la base de datos. El modelo de negocio freemium ofrece un plan Starter gratuito con 5 endpoints y 100 peticiones diarias, mientras que los planes de pago desbloquean endpoints ilimitados, mayor cuota de peticiones y retención extendida de logs.

El despliegue se realiza en infraestructura cloud moderna: el backend se conteneriza con Docker y se despliega en Railway, mientras que el frontend se aloja en Vercel con un dominio personalizado (mockagentai.com). El proyecto incluye un conjunto de pruebas de integración con JUnit 5 y H2 in-memory que validan los flujos críticos de autenticación, creación de endpoints y rate limiting.

Los resultados obtenidos demuestran la viabilidad técnica y comercial de la propuesta, validando que una arquitectura limpia con tecnologías modernas puede soportar un producto SaaS real con autenticación, autorización, control de cuotas, logs de auditoría y monetización.
"""
add_para(resumen)
add_para('Palabras clave: API mocking, SaaS, React, Spring Boot, JWT, OAuth 2.0, Stripe, agentes de inteligencia artificial, freemium.', bold=True)
doc.add_page_break()

# ÍNDICE
add_heading_custom('Índice General', 1)
indice_items = [
    "1. Introducción",
    "   1.1. Contexto y motivación",
    "   1.2. Problema identificado",
    "   1.3. Justificación del proyecto",
    "   1.4. Alcance y delimitación",
    "   1.5. Estructura de la memoria",
    "2. Objetivos",
    "   2.1. Objetivo general",
    "   2.2. Objetivos específicos",
    "3. Marco Teórico (Estado del Arte)",
    "   3.1. APIs RESTful y arquitectura cliente-servidor",
    "   3.2. API Mocking: conceptos y tipologías",
    "   3.3. Herramientas de mocking existentes",
    "   3.4. Autenticación y autorización en aplicaciones web",
    "   3.5. OAuth 2.0 e inicio de sesión social",
    "   3.6. Arquitectura de aplicaciones web modernas",
    "   3.7. Modelos de negocio SaaS y freemium",
    "   3.8. Pasarelas de pago y gestión de suscripciones",
    "4. Metodología",
    "   4.1. Metodología de desarrollo",
    "   4.2. Análisis de requisitos",
    "   4.3. Tecnologías y herramientas utilizadas",
    "   4.4. Fases del proyecto",
    "   4.5. Planificación temporal",
    "5. Implementación del Backend",
    "   5.1. Arquitectura por capas",
    "   5.2. Configuración del proyecto y dependencias",
    "   5.3. Modelo de datos y persistencia",
    "   5.4. Sistema de autenticación y autorización",
    "   5.5. Motor de mocking y simulación de APIs",
    "   5.6. Sistema de planes y rate limiting",
    "   5.7. Gestión de logs y auditoría",
    "   5.8. Seguridad y configuración de CORS",
    "6. Implementación del Frontend",
    "   6.1. Arquitectura de la aplicación",
    "   6.2. Diseño de interfaz y experiencia de usuario",
    "   6.3. Sistema de navegación y vistas",
    "   6.4. Landing page y estrategia comercial",
    "   6.5. Dashboard y gestión de endpoints",
    "   6.6. Sistema de pricing y planes",
    "   6.7. Chatbot de asistencia inteligente",
    "   6.8. Integración con el backend",
    "7. Sistema de Monetización con Stripe",
    "   7.1. Arquitectura de pagos y Stripe Checkout",
    "   7.2. Modelos de suscripción y planes de precios",
    "   7.3. Webhooks para gestión de eventos de billing",
    "   7.4. Gestión de planes y actualización de cuentas",
    "   7.5. Integración frontend-backend de pagos",
    "8. Resultados y Análisis",
    "   8.1. Funcionalidades implementadas y demostrables",
    "   8.2. Métricas de despliegue y rendimiento",
    "   8.3. Comparativa con soluciones existentes",
    "   8.4. Casos de uso y validación práctica",
    "9. Discusión",
    "   9.1. Fortalezas del proyecto",
    "   9.2. Limitaciones y amenazas identificadas",
    "   9.3. Comparación con el estado del arte",
    "   9.4. Escalabilidad y mantenimiento a largo plazo",
    "10. Conclusiones",
    "   10.1. Cumplimiento de objetivos",
    "   10.2. Impacto del proyecto y contribución",
    "   10.3. Líneas futuras de trabajo",
    "Referencias Bibliográficas",
    "Anexos",
    "   A. Configuración de seguridad del backend",
    "   B. Configuración de despliegue en producción",
    "   C. Esquema relacional de la base de datos"
]
for item in indice_items:
    add_para(item)
doc.add_page_break()

print("Generando documento...")

# 1. INTRODUCCIÓN
add_heading_custom('1. Introducción', 1)
add_heading_custom('1.1. Contexto y motivación', 2)
add_para("""La democratización de la inteligencia artificial, impulsada por modelos de lenguaje de gran escala (LLMs) como GPT-4, Claude o Gemini, ha propiciado el surgimiento de un nuevo paradigma de desarrollo software: los agentes de IA autónomos. Estos agentes son sistemas capaces de percibir su entorno, tomar decisiones y ejecutar acciones mediante el consumo de APIs externas, bases de datos y servicios de terceros. El ciclo de desarrollo de un agente de IA implica iteraciones constantes de prueba y error, donde cada petición de prueba a una API externa conlleva costes directos, dependencias de disponibilidad y limitaciones de rate limiting impuestas por los proveedores.

En este escenario, el desarrollador se enfrenta a un dilema: por un lado, necesita probar la lógica de su agente en condiciones realistas; por otro, cada ejecución de prueba consume créditos de API, expone la aplicación a cuotas estrictas y genera dependencias con servicios que pueden estar caídos o latentes. Esta problemática es particularmente aguda en el desarrollo de agentes multi-paso, donde una única ejecución puede generar entre 5 y 20 llamadas a APIs externas.

El API mocking emerge como una disciplina técnica que permite simular el comportamiento de APIs reales mediante respuestas predefinidas, eliminando la dependencia de servicios externos durante las fases de desarrollo y testing. Sin embargo, las herramientas existentes presentan carencias significativas para este nuevo paradigma: o bien son demasiado complejas para prototipado rápido (WireMock, Postman), o bien carecen de mecanismos de autenticación, planes de uso y despliegue cloud (Mockbin, Beeceptor gratuito).""")

add_heading_custom('1.2. Problema identificado', 2)
add_para("""Tras un análisis del ecosistema de desarrollo de agentes de IA y las herramientas de mocking disponibles, se identifican los siguientes problemas concretos:

1. Costes de desarrollo elevados: Cada petición de prueba a APIs de terceros (OpenAI, Stripe, SerpAPI, etc.) conlleva costes directos que dificultan la iteración ágil durante el desarrollo.

2. Rate limiting y throttling: Los proveedores de APIs imponen límites estrictos de peticiones por minuto (RPM) y por día. Durante el testing intensivo, estos límites se agotan rápidamente, bloqueando el desarrollo.

3. Imposibilidad de simular errores: No es viable ni ético forzar errores 500, timeouts o respuestas malformadas en APIs de producción para probar la robustez del agente.

4. Aislamiento de datos: En entornos de equipo, los desarrolladores necesitan entornos aislados donde sus mocks no interfieran con los de otros compañeros.

5. Falta de soluciones SaaS integrales: No existe en el mercado una solución SaaS que combine mocking de APIs, autenticación de usuarios, control de planes de uso, logs de auditoría y monetización freemium, orientada específicamente al desarrollo de agentes de IA.""")

add_heading_custom('1.3. Justificación del proyecto', 2)
add_para("""MockAgent.AI se justifica como una respuesta técnica y comercial a los problemas identificados. Desde la perspectiva académica, el proyecto integra competencias transversales del Grado en Ingeniería Informática: diseño de arquitecturas software, desarrollo backend con frameworks empresariales, desarrollo frontend con tecnologías modernas, bases de datos relacionales, seguridad informática (autenticación, autorización, validación de datos), despliegue cloud y modelos de negocio digitales.

Desde la perspectiva profesional, el proyecto demuestra la capacidad de concebir, diseñar, implementar y desplegar un producto SaaS real, con autenticación de usuarios, control de acceso basado en planes, integración de pasarelas de pago y despliegue continuo en infraestructura cloud. El dominio personalizado (mockagentai.com) y el despliegue real en producción validan la viabilidad comercial de la propuesta.""")

add_heading_custom('1.4. Alcance y delimitación', 2)
add_para("""El alcance del proyecto incluye:

- Desarrollo completo de un backend monolítico con Spring Boot.
- Desarrollo completo de un frontend SPA con React.
- Implementación de autenticación dual (JWT + Google OAuth 2.0).
- Implementación de un motor de mocking con CRUD de endpoints y respuestas JSON.
- Implementación de Instant Mock (creación sin autenticación, expiración 24h).
- Implementación de rate limiting por usuario y plan de suscripción.
- Implementación de logs de peticiones con retención diferenciada por plan.
- Integración de Stripe para suscripciones Pro y Premium.
- Despliegue real en Railway (backend) y Vercel (frontend) con dominio propio.
- Pruebas de integración para flujos críticos.

Delimitaciones y exclusiones:

- No se incluye el desarrollo de una aplicación móvil nativa (iOS/Android).
- No se implementa un sistema de notificaciones en tiempo real (WebSockets).
- No se incluye un sistema de análisis de datos avanzado (BI, dashboards analíticos).
- El chatbot es un sistema basado en reglas, no integra un LLM externo.
- No se implementa multi-tenancy avanzado ni workspaces compartidos (roadmap futuro).""")

add_heading_custom('1.5. Estructura de la memoria', 2)
add_para("""La presente memoria se organiza en diez capítulos principales. El Capítulo 1 introduce el contexto y la problemática. El Capítulo 2 establece los objetivos del proyecto. El Capítulo 3 presenta el estado del arte y los fundamentos teóricos. El Capítulo 4 describe la metodología y planificación. Los Capítulos 5 y 6 detallan la implementación técnica del backend y frontend. El Capítulo 7 aborda la monetización con Stripe. El Capítulo 8 presenta los resultados obtenidos. El Capítulo 9 discute las fortalezas y limitaciones. Finalmente, el Capítulo 10 recoge las conclusiones y líneas futuras. La memoria se completa con las referencias bibliográficas y los anexos técnicos.""")

print("Capítulo 1 completado...")

# 2. OBJETIVOS
add_heading_custom('2. Objetivos', 1)
add_heading_custom('2.1. Objetivo general', 2)
add_para("""Diseñar, desarrollar y desplegar MockAgent.AI, una plataforma SaaS de API mocking orientada al desarrollo y testing de agentes de inteligencia artificial, que permita a los usuarios crear, gestionar y consumir endpoints simulados con respuestas JSON configurables, integrando autenticación segura, control de planes de uso, monetización mediante suscripciones y despliegue en infraestructura cloud.""")

add_heading_custom('2.2. Objetivos específicos', 2)
add_para("""1. Diseñar una arquitectura backend robusta utilizando Spring Boot 3.5 con Java 17, adoptando el patrón de arquitectura por capas (controller, service, repository, DTO, exception handler) que garantice la separación de responsabilidades, la testabilidad y la mantenibilidad del código.

2. Implementar un sistema de autenticación y autorización seguro que combine autenticación basada en tokens JWT (JSON Web Tokens) con firma HMAC-SHA256 y expiración de 24 horas, hash de contraseñas mediante BCrypt, e integración de inicio de sesión social mediante OAuth 2.0 de Google con validación server-side de tokens de identidad.

3. Desarrollar un motor de mocking funcional que permita a los usuarios crear endpoints REST personalizados con métodos HTTP configurables (GET, POST, PUT, DELETE, PATCH), códigos de estado HTTP personalizables, cuerpos de respuesta JSON, y encabezados configurables. El motor debe ser capaz de recibir peticiones en rutas dinámicas y devolver las respuestas configuradas sin necesidad de código adicional.

4. Implementar un sistema de planes de uso y rate limiting que defina cuotas diferenciadas por nivel de suscripción: plan Starter (5 endpoints, 100 peticiones/día), plan Pro (endpoints ilimitados, 5.000 peticiones/día) y plan Premium (10.000 peticiones/día, logs 14 días, soporte prioritario). El sistema debe controlar tanto límites por minuto como límites diarios mediante contadores en memoria.

5. Integrar una pasarela de pagos para monetización mediante Stripe, implementando Stripe Checkout para suscripciones recurrentes (mensuales), webhooks para procesar eventos de suscripción, y actualización automática del plan del usuario en la base de datos.

6. Desarrollar un frontend moderno y responsive utilizando React 19 con Vite, TailwindCSS y Framer Motion, que incluya una landing page comercial, un dashboard de gestión de endpoints, un sistema de pricing con comparativa de planes, un chatbot de asistencia basado en reglas, y un sistema de tema claro/oscuro.

7. Implementar un sistema de logging y auditoría que registre cada petición realizada a los endpoints mock (método, ruta, código de estado, cuerpo de la petición, timestamp), aplicando retención de datos diferenciada por plan (24 horas para Starter, 7 días para Pro, 14 días para Premium).

8. Desplegar la aplicación en infraestructura cloud real mediante contenerización Docker del backend, despliegue en Railway (servicio de hosting para aplicaciones containerizadas), despliegue del frontend en Vercel (CDN edge), y configuración de un dominio personalizado (mockagentai.com) con DNS y certificados SSL.

9. Diseñar y ejecutar un conjunto de pruebas de integración utilizando JUnit 5, Spring Boot Test y H2 in-memory, que validen los flujos críticos del sistema: registro de usuarios, autenticación, creación de endpoints, límite de endpoints por plan, respuesta de mocks y eliminación de recursos.""")

print("Capítulo 2 completado...")

# Guardar progreso intermedio
output_path = 'C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\memoria-tfg.docx'
doc.save(output_path)
print(f"Documento guardado parcialmente en: {output_path}")
