import sys
sys.path.append('C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\venv\\Lib\\site-packages')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Abrir documento existente
doc = Document('C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\memoria-tfg.docx')

# Configurar márgenes de página para todas las secciones
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# Función para aplicar estilos a títulos
def style_heading(paragraph, level, font_size, color=None, bold=True):
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        run.font.name = 'Calibri'

# Función para aplicar estilos a párrafos normales
def style_paragraph(paragraph, font_size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = alignment
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.name = 'Calibri'

# Aplicar estilos a todos los párrafos
for para in doc.paragraphs:
    text = para.text.strip()
    
    # Detectar si es título de capítulo (Heading 1)
    if para.style.name.startswith('Heading 1'):
        style_heading(para, 1, 18, RGBColor(0x0A, 0x0A, 0x1E), True)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Añadir espacio antes y después
        para.paragraph_format.space_before = Pt(24)
        para.paragraph_format.space_after = Pt(12)
        
    # Detectar si es sub-título (Heading 2)
    elif para.style.name.startswith('Heading 2'):
        style_heading(para, 2, 14, RGBColor(0xC9, 0xA9, 0x6E), True)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(8)
        
    # Detectar si es sub-sub-título (Heading 3)
    elif para.style.name.startswith('Heading 3'):
        style_heading(para, 3, 12, RGBColor(0x66, 0x66, 0x66), True)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        
    # Párrafos normales
    else:
        style_paragraph(para, 11, WD_ALIGN_PARAGRAPH.JUSTIFY)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.15

# Añadir numeración de páginas
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # Crear párrafo en el footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = ""
    
    # Añadir número de página
    run = footer_para.add_run("Página ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Añadir campo PAGE
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run2 = footer_para.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    run3 = footer_para.add_run(" de ")
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Añadir campo NUMPAGES
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = "NUMPAGES"
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    run4 = footer_para.add_run()
    run4._r.append(fldChar3)
    run4._r.append(instrText2)
    run4._r.append(fldChar4)
    run4.font.size = Pt(9)
    run4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Guardar documento formateado
output_path = 'C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\memoria-tfg-formatted.docx'
doc.save(output_path)
print("[OK] Documento formateado guardado en: " + output_path)
print("\nMejoras aplicadas:")
print("   - Márgenes profesionales (3cm izquierda, 2.5cm resto)")
print("   - Títulos de capítulos: 18pt, color oscuro")
print("   - Sub-títulos: 14pt, color dorado (#C9A96E)")
print("   - Párrafos: 11pt, justificados, interlineado 1.15")
print("   - Numeracion de paginas en el footer")
print("   - Espaciado consistente entre secciones")
print("\nNOTA: Recuerda reemplazar los placeholders:")
print("   [Nombre del estudiante], [Nombre del tutor], etc.")
