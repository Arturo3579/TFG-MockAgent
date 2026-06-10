import sys
sys.path.append('C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\venv\\Lib\\site-packages')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Abrir documento existente
doc = Document('C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\memoria-tfg.docx')

def add_heading_custom(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

# 3. MARCO TEÓRICO
add_heading_custom('3. Marco Teórico (Estado del Arte)', 1)
add_heading_custom('3.1. APIs RESTful y arquitectura cliente-servidor', 2)
add_para("""La arquitectura REST (Representational State Transfer), definida por Roy Fielding en su tesis doctoral de 2000, se ha consolidado como el paradigma dominante para el diseño de APIs web. REST propone un conjunto de restricciones arquitectónicas que garantizan la escalabilidad, la fiabilidad y la evolución de los sistemas distribuidos: stateless entre peticiones, interfaz uniforme (identificación de recursos, manipulación mediante representaciones, mensajes autodescriptivos), y separación de cliente y servidor.

En el contexto de MockAgent.AI, cada endpoint mock es un recurso REST identificado por una URI única (ej. /mock/users), accesible mediante métodos HTTP semánticos (GET para lectura, POST para creación, PUT para actualización, DELETE para eliminación). El cuerpo de las respuestas se representa en formato JSON (JavaScript Object Notation), que se ha convertido en el estándar de facto para el intercambio de datos en APIs modernas debido a su legibilidad, ligereza y compatibilidad universal.""")

add_heading_custom('3.2. API Mocking: conceptos y tipologías', 2)
add_para("""El API mocking es la práctica de simular el comportamiento de una API real mediante un servicio intermedio que responde con datos predefinidos. Existen tres tipologías principales:

1. Client-side mocking: Se implementa en el código del cliente (frontend o aplicación móvil) mediante bibliotecas como Mock Service Worker (MSW) o interceptores de HTTP. Útil para testing unitario pero no sirve para testing de integración ni para desarrollo de backend.

2. Server-side mocking: Se implementa en un servidor independiente que expone endpoints reales y responde con datos configurados. Es el enfoque de MockAgent.AI, ya que permite que cualquier cliente (cURL, Postman, agentes de IA) consuma los endpoints sin modificar el código del cliente.

3. Contract-based mocking: Se basa en definir contratos formales (OpenAPI, Pact, GraphQL schemas) y generar mocks automáticamente a partir de esos contratos. Herramientas como WireMock o Prism siguen este enfoque.

MockAgent.AI adopta el enfoque server-side con una capa de configuración manual, permitiendo al usuario definir exactamente la respuesta que desea, incluyendo códigos de error, latencia simulada (futura) y cuerpos JSON complejos.""")

add_heading_custom('3.3. Herramientas de mocking existentes', 2)
add_para("""El ecosistema de herramientas de mocking es amplio y maduro. A continuación se analizan las principales:

Postman Mock Server: Integrado en la suite de Postman, permite generar mocks a partir de colecciones de API. Requiere colecciones predefinidas y su plan gratuito es limitado. No está orientado a agentes de IA ni incluye autenticación de usuarios propia.

WireMock: Biblioteca Java de código abierto muy potente para testing. Requiere conocimientos técnicos avanzados, configuración mediante JSON o Java, y no ofrece una interfaz web para usuarios no técnicos. Su despliegue es autogestionado.

MockServer: Herramienta genérica que permite crear mocks mediante expectativas programáticas. Similar a WireMock en complejidad, orientada a entornos de testing automatizado, no a producto SaaS.

Beeceptor: Servicio SaaS que permite interceptar y simular endpoints. Ofrece una interfaz web amigable pero carece de autenticación robusta, planes de uso diferenciados y sistema de logs avanzado. Su modelo gratuito es muy restrictivo.

Mockbin: Servicio simple de creación de bins que registran peticiones. No permite respuestas configurables de forma nativa ni autenticación de usuarios.

MockAgent.AI se posiciona en el espacio entre Beeceptor (simplicidad SaaS) y WireMock (potencia técnica), añadiendo autenticación JWT, OAuth 2.0, planes de uso, logs de auditoría, rate limiting y monetización Stripe, todo ello orientado al nicho de agentes de IA.""")

add_heading_custom('3.4. Autenticación y autorización en aplicaciones web', 2)
add_para("""La autenticación es el proceso de verificar la identidad de un usuario, mientras que la autorización determina qué recursos puede acceder una vez autenticado. En aplicaciones web modernas, el estándar dominante es el uso de tokens JWT (JSON Web Tokens, RFC 7519).

Un JWT es un token compacto y autónomo que transmite información entre partes como un objeto JSON firmado digitalmente. Consta de tres partes: header (algoritmo de firma), payload (claims como email, plan, expiración) y signature (firma criptográfica). La firma HMAC-SHA256 garantiza la integridad del token: si un atacante modifica el payload, la firma ya no coincide y el token se rechaza.

En MockAgent.AI, el backend genera un JWT tras el login exitoso (tanto por email/password como por Google OAuth), lo devuelve al frontend, y el frontend lo almacena en localStorage o sessionStorage según la preferencia del usuario (Recuérdame). En cada petición posterior, el frontend incluye el JWT en el header Authorization: Bearer <token>. El backend valida el JWT mediante un filtro (JwtAuthFilter) que intercepta todas las peticiones y verifica la firma y la expiración.

Para el hash de contraseñas, se utiliza BCrypt, un algoritmo de hashing adaptativo que incluye un salt automático y un factor de coste configurable. Spring Security proporciona BCryptPasswordEncoder que se configura como bean en la aplicación.""")

add_heading_custom('3.5. OAuth 2.0 e inicio de sesión social', 2)
add_para("""OAuth 2.0 es un protocolo de autorización que permite a las aplicaciones obtener acceso limitado a cuentas de usuario en servicios de terceros (Google, GitHub, Facebook) sin necesidad de conocer la contraseña del usuario. En el contexto de inicio de sesión social, el flujo utilizado es el Authorization Code Flow con PKCE (Proof Key for Code Exchange), o bien el flujo simplificado mediante Google Identity Services (ID tokens).

MockAgent.AI implementa el flujo de Google Identity Services mediante la librería @react-oauth/google en el frontend. Cuando el usuario pulsa Sign in with Google, Google devuelve un idToken (JWT firmado por Google) al frontend. El frontend envía este token al backend (POST /api/auth/google). El backend utiliza la librería google-api-client de Google para verificar la firma del token, extraer el email y el googleId (sub), y crear o vincular la cuenta de usuario. El backend genera entonces su propio JWT interno y lo devuelve al frontend, manteniendo la coherencia del sistema de autenticación independientemente del método de login.""")

add_heading_custom('3.6. Arquitectura de aplicaciones web modernas', 2)
add_para("""Las aplicaciones web modernas siguen una arquitectura de separación de frontend y backend. El frontend es una Single Page Application (SPA) que se ejecuta en el navegador del cliente, consume datos mediante peticiones asíncronas (AJAX) al backend, y actualiza la interfaz dinámicamente sin recargar la página completa. El backend expone una API RESTful que actúa como capa de negocio y persistencia. La comunicación entre frontend y backend se realiza mediante HTTP/HTTPS y JSON.

En el frontend, el ecosistema React domina el mercado. React 19 introduce mejoras en el manejo de estado, el servidor de componentes y la concurrencia. Vite ha reemplazado a Create React App como herramienta de build estándar debido a su velocidad de compilación. TailwindCSS ha popularizado el enfoque utility-first, donde los estilos se aplican mediante clases atómicas directamente en el JSX, eliminando la necesidad de escribir CSS personalizado.

En el backend, Spring Boot es el framework de facto para aplicaciones Java empresariales. Su modelo de inversión de control (IoC) y inyección de dependencias (DI) facilita la configuración y el testing. Spring Data JPA abstrae el acceso a bases de datos relacionales mediante repositorios que generan consultas SQL automáticamente.""")

add_heading_custom('3.7. Modelos de negocio SaaS y freemium', 2)
add_para("""El modelo SaaS (Software as a Service) consiste en ofrecer software alojado en la nube, accesible mediante suscripción, sin necesidad de instalación local. El modelo freemium es una variante del SaaS donde se ofrece una versión gratuita con funcionalidades limitadas y versiones de pago con funcionalidades avanzadas.

Los componentes clave de un SaaS freemium son: Onboarding (proceso de registro fluido), Plan differentiation (cuotas claras que incentiven la conversión), Payment gateway (pasarela segura), Billing management (gestión de suscripciones y facturas), y Usage-based alerts (notificaciones cuando el usuario se acerca a su límite).

MockAgent.AI adopta el modelo freemium con tres planes: Starter (gratuito), Pro ($4.99/mes) y Premium ($7.99/mes). La conversión se incentiva mediante límites de endpoints y peticiones, así como retención de logs.""")

add_heading_custom('3.8. Pasarelas de pago y gestión de suscripciones', 2)
add_para("""Stripe es la pasarela de pagos líder para SaaS y startups. Su API RESTful permite crear sesiones de checkout, gestionar suscripciones, procesar webhooks y generar facturas. Los componentes clave utilizados en MockAgent.AI son:

Stripe Checkout: Página de pago alojada por Stripe que recopila los datos de la tarjeta de forma segura (PCI DSS compliant), reduciendo la responsabilidad de cumplimiento del desarrollador.

Stripe Billing: Sistema de suscripciones recurrentes que maneja ciclos de facturación, pruebas gratuitas, cancelaciones y renovaciones.

Stripe Webhooks: Notificaciones HTTP que Stripe envía a la aplicación del desarrollador cuando ocurren eventos (pago exitoso, suscripción cancelada, tarjeta rechazada). El backend debe validar la firma del webhook para evitar spoofing.

Stripe Customer Portal: Interfaz prefabricada donde los usuarios pueden gestionar su método de pago, ver facturas y cancelar suscripciones.""")

print("Capítulo 3 completado...")

# Guardar
output_path = 'C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\memoria-tfg.docx'
doc.save(output_path)
print(f"Documento actualizado en: {output_path}")
