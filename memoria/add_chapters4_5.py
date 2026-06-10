import sys
sys.path.append('C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\venv\\Lib\\site-packages')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\memoria-tfg.docx')

def add_heading_custom(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

# 4. METODOLOGÍA
add_heading_custom('4. Metodología', 1)
add_heading_custom('4.1. Metodología de desarrollo', 2)
add_para("""El proyecto se ha desarrollado siguiendo una adaptación de la metodología ágil Scrum, con sprints de duración variable (1-2 semanas) y un enfoque iterativo incremental. Dado que el proyecto es un trabajo individual de fin de grado, no se han realizado ceremonias formales de Scrum (daily standups, sprint planning con equipo), pero sí se ha aplicado el espíritu iterativo: desarrollo de funcionalidades completas en ciclos cortos, integración continua mediante Git, despliegue frecuente para validación en producción, y adaptación de requisitos según surgían necesidades durante la implementación.

La gestión del código se ha realizado mediante Git, con un repositorio único en GitHub que contiene ambos proyectos (backend y frontend) en directorios separados. El flujo de trabajo es sencillo: rama main como principal, commits frecuentes con mensajes descriptivos (prefijos feat:, fix:, ui:, chore:), y push directo a main dado que es un proyecto individual. Railway y Vercel están configurados con auto-deploy desde la rama main, por lo que cada push desencadena un nuevo despliegue en producción.""")

add_heading_custom('4.2. Análisis de requisitos', 2)
add_para("""El análisis de requisitos se ha realizado mediante técnicas de brainstorming, análisis de competidores y definición de casos de uso. Los requisitos se han clasificado en funcionales y no funcionales.

Requisitos funcionales (RF):
- RF1: El sistema debe permitir el registro de usuarios mediante email y contraseña.
- RF2: El sistema debe permitir el inicio de sesión mediante email/contraseña y Google OAuth.
- RF3: El sistema debe permitir la creación de endpoints mock con path, método, status y body JSON.
- RF4: El sistema debe permitir la edición y eliminación de endpoints propios.
- RF5: El sistema debe responder a peticiones en /mock/{path} con la configuración del endpoint.
- RF6: El sistema debe permitir la creación de Instant Mocks sin autenticación (expiración 24h).
- RF7: El sistema debe registrar logs de cada petición a los endpoints.
- RF8: El sistema debe aplicar rate limiting por usuario (minuto y diario).
- RF9: El sistema debe soportar tres planes: Starter, Pro, Premium.
- RF10: El sistema debe integrar Stripe para suscripciones de pago.
- RF11: El sistema debe permitir upgrade/downgrade de plan.
- RF12: El sistema debe ofrecer un dashboard de gestión y un landing page.
- RF13: El sistema debe incluir un chatbot de asistencia basado en reglas.

Requisitos no funcionales (RNF):
- RNF1: Seguridad: contraseñas hasheadas con BCrypt, tokens JWT firmados, validación de entrada, CORS configurado.
- RNF2: Rendimiento: tiempo de respuesta < 200ms para mocks, < 500ms para operaciones CRUD.
- RNF3: Disponibilidad: despliegue en cloud con uptime objetivo del 99.9%.
- RNF4: Escalabilidad: arquitectura stateless que permite escalado horizontal.
- RNF5: Usabilidad: interfaz responsive, tema claro/oscuro, flujo de registro < 2 minutos.
- RNF6: Mantenibilidad: código limpio, arquitectura por capas, tests de integración.""")

add_heading_custom('4.3. Tecnologías y herramientas utilizadas', 2)
add_para("""Backend:
- Lenguaje: Java 17 (LTS, soporte extendido, mejoras en pattern matching y records)
- Framework: Spring Boot 3.5.14 (inyección de dependencias, auto-configuración, starters)
- Seguridad: Spring Security (filtros, BCrypt, stateless sessions)
- Persistencia: Spring Data JPA (Hibernate 6.x), MySQL 8.0 driver
- Token JWT: JJWT 0.12.6 (API moderna, inmutable, compatible con Spring Boot 3)
- Validación: Jakarta Validation (@NotBlank, @Email, @Size)
- Build: Maven 3.9 (gestión de dependencias, plugins, Dockerfile multi-stage)
- OAuth: Google API Client 2.2.0 (verificación de ID tokens)
- Pagos: Stripe Java SDK 24.0.0 (creación de sesiones, webhooks)
- Testing: JUnit 5, Spring Boot Test, TestRestTemplate, H2 database
- Contenerización: Docker, eclipse-temurin:17-jdk-alpine / 17-jre-alpine

Frontend:
- Framework: React 19.2.6 (hooks, concurrent features, compiler)
- Build tool: Vite 8.0.12 (dev server con HMR, build optimizado con Rollup)
- Estilos: TailwindCSS 4.3.0 (utility-first, CSS variables para theming)
- Animaciones: Framer Motion 12.40.0 (transiciones, gestos, layout animations)
- Iconos: Lucide React 1.17.0 (iconos SVG, tree-shakeable)
- HTTP client: Axios 1.17.0 (interceptores, configuración global)
- OAuth: @react-oauth/google 0.13.5 (GoogleLogin component, context provider)
- Router: No se utiliza react-router; navegación mediante estado vistaActual (SPA condicional)
- Linter: ESLint 10.3.0

Infraestructura y DevOps:
- Control de versiones: Git, GitHub (Arturo3579/TFG-MockAgent)
- Backend hosting: Railway (deploy automático desde GitHub, variables de entorno, health checks)
- Frontend hosting: Vercel (deploy automático, edge network, SPA rewrite rules)
- Base de datos: TiDB Cloud (MySQL compatible, serverless, auto-scaling)
- Dominio: Namecheap (mockagentai.com, CNAME www, A record @, verificado en Vercel)
- SSL: Vercel y Railway proporcionan certificados HTTPS automáticos

Herramientas de desarrollo:
- IDE: VS Code (con extensiones para Java, Spring Boot, React, Tailwind, Git)
- API Testing: Postman, cURL (validación de endpoints en local y producción)
- Navegadores: Chrome, Firefox, Safari (validación de responsive y compatibilidad)
- Documentación: Markdown (README.md, DEPLOY.md, CHANGELOG.md)""")

add_heading_custom('4.4. Fases del proyecto', 2)
add_para("""El proyecto se ha dividido en las siguientes fases secuenciales:

Fase 1: Análisis y planificación (Semanas 1-2)
- Definición del problema y del alcance.
- Análisis de competidores (Beeceptor, Postman, WireMock).
- Selección de tecnologías y arquitectura.
- Creación de repositorio y configuración inicial.

Fase 2: Diseño del backend (Semanas 3-4)
- Configuración de Spring Boot con Spring Initializr.
- Diseño del modelo de datos (entidades, relaciones, enums).
- Implementación de repositorios JPA.
- Configuración de seguridad (SecurityConfig, JWT, BCrypt).
- Implementación de AuthService y AuthController.

Fase 3: Implementación del motor de mocking (Semanas 5-6)
- Desarrollo de EndpointService y EndpointController.
- Implementación de MockController (motor de matching path+method).
- Desarrollo de InstantController (mock sin autenticación).
- Implementación de validación JSON y control de planes.

Fase 4: Implementación del frontend (Semanas 7-10)
- Configuración de Vite + React + Tailwind.
- Desarrollo del sistema de estado (vistaActual, token, tema).
- Implementación de Landing Page.
- Implementación de Login/Signup.
- Implementación de Dashboard.
- Implementación de Pricing, Docs, Blog, Profile.
- Integración con backend mediante Axios.

Fase 5: Sistema de monetización (Semanas 11-12)
- Configuración de cuenta Stripe y productos (Pro, Premium).
- Implementación de Stripe Checkout en backend.
- Desarrollo de webhook handler para eventos de billing.
- Integración de botones de upgrade en frontend.
- Implementación de customer portal.

Fase 6: Testing y despliegue (Semanas 13-14)
- Escritura de tests de integración (JUnit 5, H2).
- Contenerización con Docker (Dockerfile multi-stage).
- Despliegue en Railway (backend) y Vercel (frontend).
- Configuración de dominio custom y DNS.
- Validación de flujos críticos en producción.

Fase 7: Pulido y documentación (Semanas 15-16)
- Mejora de responsive móvil (navbar hamburger, chat widget).
- Mejora del chatbot IA (memoria de contexto, follow-ups).
- Implementación de Google Sign-In.
- Actualización de README y DEPLOY.md.
- Redacción de la memoria de TFG.""")

add_heading_custom('4.5. Planificación temporal', 2)
add_para("""La planificación temporal se ha gestionado de forma flexible, sin un diagrama de Gantt formal, pero con hitos semanales claros. La duración total ha sido de aproximadamente 16 semanas (4 meses), con una dedicación part-time de unas 20-25 horas semanales. Los hitos principales fueron:

- Hito 1 (Semana 2): Backend con autenticación JWT funcional.
- Hito 2 (Semana 4): CRUD de endpoints y motor de mocking operativo.
- Hito 3 (Semana 6): Frontend con login, signup y dashboard funcional.
- Hito 4 (Semana 8): Landing page, pricing y docs desplegados en Vercel.
- Hito 5 (Semana 10): Rate limiting, logs, Instant Mock y control de planes.
- Hito 6 (Semana 12): Integración de Stripe Checkout y webhooks.
- Hito 7 (Semana 14): Backend desplegado en Railway, frontend en Vercel, dominio configurado.
- Hito 8 (Semana 16): Google Sign-In, chatbot mejorado, responsive final, memoria completa.""")

print("Capítulo 4 completado...")

# 5. IMPLEMENTACIÓN DEL BACKEND
add_heading_custom('5. Implementación del Backend', 1)
add_heading_custom('5.1. Arquitectura por capas', 2)
add_para("""El backend de MockAgent.AI sigue una arquitectura por capas clásica, organizada en paquetes que reflejan la separación de responsabilidades. Esta arquitectura es monolítica pero modular, lo que facilita el testing y el mantenimiento. Las capas son:

1. Capa de presentación (Controllers): Recibe las peticiones HTTP, valida los parámetros de entrada, delega la lógica a los servicios y devuelve las respuestas HTTP. Los controladores son stateless y no contienen lógica de negocio.

2. Capa de negocio (Services): Contiene la lógica de negocio, las reglas de validación, las operaciones de transformación de datos y las llamadas a repositorios. Los servicios son anotados con @Service y se inyectan en los controladores mediante inyección de dependencias de Spring.

3. Capa de acceso a datos (Repositories): Interfaces que extienden JpaRepository o CrudRepository. Spring Data JPA genera automáticamente las implementaciones en tiempo de ejecución, proporcionando métodos CRUD básicos y consultas personalizadas mediante derivación de nombres de método (ej. findByEmail, findByGoogleId).

4. Capa de modelo (Models/Entities): Clases POJO anotadas con JPA (@Entity, @Table, @Id, @GeneratedValue) que mapean objetos Java a tablas de la base de datos relacional.

5. Capa de DTOs (Data Transfer Objects): Objetos inmutables (o casi inmutables) que transportan datos entre capas y entre frontend y backend. Los DTOs desacoplan la API pública de la representación interna de la base de datos.

6. Capa de excepciones y seguridad: Incluye excepciones personalizadas (UserNotFoundException, InvalidCredentialsException, RateLimitExceededException), un manejador global de excepciones (GlobalExceptionHandler) que devuelve respuestas JSON estandarizadas con código HTTP apropiado, y los componentes de seguridad (SecurityConfig, JwtAuthFilter, JwtUtil).

7. Capa de utilidades: Validadores personalizados (EmailValidator, PasswordValidator) que encapsulan reglas de negocio reutilizables.

Esta arquitectura garantiza que una modificación en la capa de persistencia (por ejemplo, cambiar de MySQL a PostgreSQL) no afecte a la capa de negocio, y que un cambio en la API REST (por ejemplo, añadir un campo al DTO) no requiera modificar la base de datos si no es necesario.""")

add_heading_custom('5.2. Configuración del proyecto y dependencias', 2)
add_para("""El proyecto backend se configura como un proyecto Maven con pom.xml como descriptor de dependencias. El parent POM es spring-boot-starter-parent versión 3.5.14, que gestiona automáticamente las versiones compatibles de las dependencias Spring.

Las dependencias principales son:
- spring-boot-starter-web: Incluye Tomcat embebido, Spring MVC, Jackson para serialización JSON.
- spring-boot-starter-data-jpa: Incluye Hibernate, Spring Data JPA, y el pool de conexiones HikariCP.
- spring-boot-starter-security: Incluye Spring Security, BCrypt, y el framework de filtros.
- spring-boot-starter-validation: Incluye Jakarta Bean Validation (Hibernate Validator).
- jjwt-api, jjwt-impl, jjwt-jackson: Implementación de JWT con soporte para parsing JSON via Jackson.
- mysql-connector-j: Driver JDBC para MySQL 8.0.
- stripe-java: SDK oficial de Stripe para Java.
- google-api-client: Cliente oficial de Google para validar tokens OAuth.
- lombok: Reducción de boilerplate mediante anotaciones (@Getter, @Setter, @Builder).
- h2: Base de datos en memoria para tests (scope test).

El archivo application.properties configura:
- spring.datasource.url: URL JDBC de la base de datos (MySQL en producción, variable de entorno DATABASE_URL en Railway).
- spring.jpa.hibernate.ddl-auto=update: Estrategia de generación de esquema (crea/actualiza tablas automáticamente).
- spring.jpa.properties.hibernate.dialect=org.hibernate.dialect.MySQLDialect: Dialecto SQL para MySQL.
- jwt.secret: Clave secreta para firmar JWT (variable de entorno JWT_SECRET).
- jwt.expiration=86400000: Expiración de tokens en milisegundos (24 horas).
- google.client.id: Client ID de Google OAuth (variable de entorno GOOGLE_CLIENT_ID).
- stripe.secret.key, stripe.webhook.secret, stripe.price.pro.monthly, stripe.price.premium.monthly: Configuración de Stripe (variables de entorno).
- server.port=${PORT:8080}: Puerto dinámico para Railway (usa variable PORT si existe, 8080 por defecto).""")

add_heading_custom('5.3. Modelo de datos y persistencia', 2)
add_para("""El modelo de datos se compone de tres entidades principales, diseñadas para satisfacer los requisitos de autenticación, mocking y auditoría.

Entidad User (tabla users):
- id: Long, clave primaria, autoincremental.
- email: String, unique, not null. Correo electrónico del usuario.
- password: String, not null. Hash BCrypt de la contraseña. Para usuarios de Google OAuth, se genera un hash dummy de una UUID aleatoria para mantener la restricción not null sin exponer la cuenta.
- googleId: String, unique, nullable. Identificador de Google (sub) para usuarios que se registran mediante OAuth. Permite vincular una cuenta de Google a una cuenta existente por email.
- plan: Enum PlanType, not null, default STARTER. Define el nivel de suscripción.

Entidad MockEndpoint (tabla mock_endpoints):
- id: Long, clave primaria.
- user: Relación Many-to-One con User (nullable). Si es null, el endpoint es público (Instant Mock). Si tiene usuario, es privado y solo accesible por el propietario.
- path: String, not null. Ruta del endpoint (ej. users, products/123).
- method: String, not null. Método HTTP (GET, POST, PUT, DELETE, PATCH).
- status: Integer, default 200. Código de estado HTTP de la respuesta.
- responseBody: String, tipo TEXT. Cuerpo de la respuesta en formato JSON.
- isDemo: Boolean, default false. Indica si es un endpoint de demostración.
- createdAt: Timestamp. Fecha de creación.

Entidad RequestLog (tabla request_logs):
- id: Long, clave primaria.
- endpoint: Relación Many-to-One con MockEndpoint. Puede ser null para peticiones a Instant Mock.
- user: Relación Many-to-One con User. Puede ser null para peticiones no autenticadas.
- timestamp: Timestamp. Momento de la petición.
- method: String. Método HTTP recibido.
- path: String. Ruta recibida.
- statusCode: Integer. Código de estado devuelto.
- requestBody: String, tipo TEXT. Cuerpo de la petición recibida.

Enum PlanType:
- STARTER: Plan gratuito (5 endpoints, 100 peticiones/día, logs 24h).
- PRO: Plan de pago ($4.99/mes, endpoints ilimitados, 5.000 peticiones/día, logs 7 días).
- PREMIUM: Plan de pago ($7.99/mes, 10.000 peticiones/día, logs 14 días, soporte prioritario).
- ENTERPRISE: Plan reservado para futuros clientes corporativos.

La estrategia de generación de esquema es ddl-auto=update, gestionada por Hibernate. No se utilizan migraciones formales (Flyway/Liquibase) dado el contexto académico y la naturaleza evolutiva del proyecto, aunque esta es una limitación reconocida para producción a gran escala.""")

add_heading_custom('5.4. Sistema de autenticación y autorización', 2)
add_para("""El sistema de autenticación es dual, permitiendo dos vías de acceso independientes que convergen en el mismo sistema de autorización basado en JWT.

Vía 1: Email y contraseña (JWT clásico)
1. El usuario envía email y contraseña al endpoint POST /api/auth/signup o POST /api/auth/login.
2. AuthController recibe el DTO AuthRequest (validado con @Valid, @NotBlank, @Email, @Size).
3. Para signup: AuthService verifica que el email no exista (existsByEmail), valida el formato del email (EmailValidator), valida la fortaleza de la contraseña (PasswordValidator: mínimo 8 caracteres, 1 mayúscula, 1 minúscula, 1 número, 1 especial), hashea la contraseña con BCryptPasswordEncoder.encode(), crea el usuario con PlanType.STARTER, y lo persiste.
4. Para login: AuthService busca el usuario por email, compara la contraseña proporcionada con el hash almacenado mediante passwordEncoder.matches(), y lanza InvalidCredentialsException si no coinciden.
5. Tras la autenticación exitosa (tanto en signup como en login), AuthService genera un JWT invocando JwtUtil.generateToken(email, plan.name()).
6. JwtUtil utiliza el algoritmo HMAC-SHA256 (HS256) para firmar el token. El payload incluye: sub (email), plan (nombre del plan), iat (issued at), exp (expiración, 24h). La firma se calcula con la clave secreta configurada en jwt.secret.
7. El token se devuelve al frontend dentro de un AuthResponse (DTO que contiene token, email, plan).

Vía 2: Google OAuth 2.0
1. El frontend renderiza el componente GoogleLogin de @react-oauth/google, envuelto en GoogleOAuthProvider con el clientId.
2. El usuario selecciona su cuenta de Google y el frontend recibe un credentialResponse con un idToken (JWT firmado por Google).
3. El frontend envía POST /api/auth/google con el DTO GoogleAuthRequest (campo idToken).
4. AuthController delega a AuthService.authenticateWithGoogle(idToken).
5. AuthService construye un GoogleIdTokenVerifier con la librería google-api-client, configurando el googleClientId como audience válida.
6. El verificador descarga las claves públicas de Google, valida la firma del token, verifica que el aud coincida con el clientId, y extrae el payload.
7. Del payload se extraen: email (claim email) y googleId (claim sub).
8. El sistema busca un usuario existente por googleId o por email. Si no existe, crea uno nuevo con googleId y una contraseña dummy (hash de UUID). Si existe por email pero sin googleId, vincula el googleId.
9. Genera un JWT interno igual que en la vía clásica y lo devuelve.

Autorización mediante JWT Filter
1. JwtAuthFilter extiende OncePerRequestFilter e intercepta todas las peticiones HTTP.
2. Extrae el header Authorization: Bearer <token>. Si no existe o no comienza con Bearer, continúa la cadena de filtros (petición pública).
3. Si existe, extrae el token, invoca JwtUtil.validateToken(token) para verificar la firma y la expiración.
4. Si el token es válido, extrae el email del claim sub y busca el usuario en la base de datos.
5. Construye un UsernamePasswordAuthenticationToken con el email y una lista vacía de authorities (no se implementa RBAC, solo control de plan basado en claim), y lo establece en el SecurityContextHolder.
6. SecurityConfig configura sessionManagement().sessionCreationPolicy(SessionCreationPolicy.STATELESS) para garantizar que Spring Security no cree sesiones HTTP server-side.
7. SecurityConfig también deshabilita CSRF (csrf.disable()) porque la API es stateless y no utiliza formularios HTML tradicionales.
8. Las rutas públicas (/, /api/health, /api/auth/**, /api/instant/**, /api/stripe/webhook, /mock/**, OPTIONS) se permiten sin autenticación mediante authorizeHttpRequests().requestMatchers(...).permitAll().""")

add_heading_custom('5.5. Motor de mocking y simulación de APIs', 2)
add_para("""El motor de mocking es el núcleo funcional de la aplicación. Consta de dos componentes: el CRUD de endpoints (privados, con autenticación) y el motor de resolución de peticiones (público o privado, según el tipo de endpoint).

CRUD de endpoints (EndpointController y EndpointService)
- POST /api/admin/endpoints: Crea un endpoint. Recibe EndpointRequest (path, method, status, responseBody). Valida que el body sea JSON válido (isValidJson). Valida que el usuario no haya excedido el límite de endpoints de su plan (countByUser >= getEndpointLimit(plan)). Valida que el path no contenga caracteres inválidos (/). Persiste el endpoint asociado al usuario autenticado.
- GET /api/admin/endpoints: Lista todos los endpoints del usuario autenticado (extraído del JWT). Soporta búsqueda por nombre de path mediante parámetro search.
- PUT /api/admin/endpoints/{id}: Actualiza un endpoint existente. Verifica que el endpoint pertenezca al usuario autenticado (evita que un usuario modifique endpoints de otro).
- DELETE /api/admin/endpoints/{id}: Elimina un endpoint propio. También elimina los logs asociados en cascada (o los deja como huérfanos, según configuración JPA).

Motor de resolución (MockController)
- RequestMapping("/mock/**"): Captura todas las peticiones que comienzan con /mock/.
- Extrae el path solicitado (request.getRequestURI().substring(6) para quitar /mock/).- Extrae el método HTTP (request.getMethod()).
- Busca en la base de datos un MockEndpoint donde path coincida y method coincida (ignorando mayúsculas/minúsculas).
- Si no encuentra, devuelve 404.
- Si encuentra, registra la petición en RequestLog (asíncrono o síncrono, según implementación).
- Devuelve una ResponseEntity con el statusCode configurado y el responseBody como JSON. Añade header Content-Type: application/json.
- Si el endpoint tiene user (privado), verifica que el JWT sea válido (aunque actualmente el motor de mocks es público por diseño; los endpoints privados se filtran en el dashboard, no en el motor de consumo).

Instant Mock (InstantController y InstantService)
- POST /api/instant/create: Recibe InstantRequest (name, method, status, response). Crea un MockEndpoint con user=null, isDemo=false, y una fecha de expiración implícita de 24 horas.
- GET /api/instant/{name}: Resuelve el mock por nombre. Verifica que no haya expirado (24h desde creación).
- @Scheduled(cron = "0 0 * * * *"): Tarea programada cada hora que elimina los Instant Mocks creados hace más de 24 horas.
- GET /api/instant/list: Lista todos los Instant Mocks activos (para descubrimiento público).""")

add_heading_custom('5.6. Sistema de planes y rate limiting', 2)
add_para("""El sistema de planes se implementa mediante el enum PlanType y el servicio RateLimitService. La lógica de control de planes reside en EndpointService (para límites de endpoints) y RateLimitService (para límites de peticiones).

Control de endpoints por plan:
- getEndpointLimit(PlanType plan): Devuelve el límite según el plan (STARTER=5, PRO=Integer.MAX_VALUE, PREMIUM=Integer.MAX_VALUE).
- countByUser(User user): Cuenta los endpoints del usuario.
- Si count >= limit y el método es CREATE, lanza PlanLimitExceededException (403 Forbidden).

Rate limiting:
- RateLimitService utiliza ConcurrentHashMap para almacenar ventanas de tiempo en memoria. Esta implementación es suficiente para un SaaS en fase inicial con un único servidor, aunque no escala a arquitecturas multi-nodo (que requerirían Redis).
- Límite por minuto: Se registra la marca de tiempo de la última petición del usuario. Si el usuario excede el RPM (Requests Per Minute) de su plan, se lanza RateLimitExceededException (429 Too Many Requests).
- Límite diario: Se mantiene un contador diario por usuario. Si excede el límite diario (100 para Starter, 5000 para Pro, 10000 para Premium), se lanza RateLimitExceededException.
- Reset: El contador diario se resetea a las 00:00 UTC (o según lógica de ventana deslizante, según implementación exacta).

Upgrade/Downgrade de plan:
- POST /api/auth/upgrade: Recibe plan (String). Actualiza el campo plan del usuario autenticado.
- En la implementación actual, este endpoint es una simulación para testing. En la versión con Stripe, el plan se actualiza automáticamente mediante el webhook de Stripe (invoice.payment_succeeded para activación, customer.subscription.deleted para cancelación).""")

add_heading_custom('5.7. Gestión de logs y auditoría', 2)
add_para("""Cada petición recibida por el motor de mocking (MockController) se registra en la entidad RequestLog. Los logs son esenciales para debugging, análisis de uso y auditoría.

Campos registrados:
- Timestamp exacto de la petición.
- Método HTTP (GET, POST, etc.).
- Path completo solicitado.
- Código de estado devuelto.
- Cuerpo de la petición (request body), almacenado como TEXT para peticiones grandes.
- Relación con el endpoint (si aplica) y con el usuario (si la petición está autenticada).

Retención de logs por plan:
- RequestLogService implementa un método de purga periódica (o manual) que elimina logs antiguos según el plan del usuario:
  - STARTER: logs de más de 24 horas.
  - PRO: logs de más de 7 días.
  - PREMIUM: logs de más de 14 días.
- Esta purga se puede ejecutar mediante una tarea @Scheduled diaria, o manualmente desde el dashboard (botón Limpiar logs antiguos).

Visualización de logs:
- GET /api/admin/logs: Devuelve los logs del usuario autenticado, ordenados por timestamp descendente.
- El frontend muestra los logs en una tabla con paginación (o scroll infinito), mostrando método, path, status, timestamp y un botón para ver el cuerpo completo de la petición.""")

add_heading_custom('5.8. Seguridad y configuración de CORS', 2)
add_para("""La configuración de seguridad se centraliza en SecurityConfig.java, que define la cadena de filtros de Spring Security y las políticas de acceso.

Configuración CORS (Cross-Origin Resource Sharing):
- El backend permite peticiones desde orígenes específicos, no desde cualquier origen (*):
  - http://localhost:5173, http://localhost:3000, http://localhost:4173 (desarrollo local con Vite).
  - https://tfg-mockagent.vercel.app (despliegue anterior).
  - https://mockagent-ai.vercel.app (despliegue alternativo).
  - https://mockagentai.com y https://www.mockagentai.com (dominio actual en producción).
- Métodos permitidos: GET, POST, PUT, DELETE, OPTIONS.
- Headers permitidos: Authorization, Content-Type, X-Requested-With.
- allowCredentials(true): Permite el envío de cookies y headers de autorización en peticiones CORS.
- maxAge(3600): Cache de preflight de 1 hora.

Configuración de seguridad HTTP:
- csrf(AbstractHttpConfigurer::disable): Deshabilita CSRF porque la API es stateless (no hay sesiones server-side ni formularios tradicionales). El token JWT se envía en el header, no en cookies, por lo que CSRF no es un vector de ataque relevante.
- sessionManagement(session -> session.sessionCreationPolicy(SessionCreationPolicy.STATELESS)): Garantiza que Spring Security no cree sesiones HTTP. Cada petición es independiente y se autentica mediante el JWT.
- authorizeHttpRequests(...): Define las reglas de autorización:
  - requestMatchers("/", "/api/health", "/api/auth/**", "/api/instant/**", "/api/stripe/webhook", "/mock/**", "OPTIONS /**").permitAll(): Rutas públicas.
  - anyRequest().authenticated(): Cualquier otra ruta requiere autenticación válida.
- addFilterBefore(jwtAuthFilter, UsernamePasswordAuthenticationFilter.class): Inserta el filtro JWT antes del filtro de autenticación por usuario/contraseña de Spring Security.

Validación de entrada:
- Todos los endpoints que reciben datos del cliente utilizan DTOs anotados con Jakarta Validation:
  - @NotBlank: Campos obligatorios no vacíos.
  - @Email: Formato de email válido.
  - @Size: Longitud mínima/máxima de strings.
- GlobalExceptionHandler captura MethodArgumentNotValidException y devuelve un ApiError con los campos inválidos y mensajes descriptivos.
- EmailValidator implementa una regex robusta para validar emails y una lista de dominios temporales desechables (tempmail, 10minutemail, etc.) para evitar registros de spam.
- PasswordValidator valida la fortaleza de la contraseña (mínimo 8 caracteres, 1 mayúscula, 1 minúscula, 1 número, 1 especial).

Protección contra SQL Injection:
- Spring Data JPA genera consultas parametrizadas automáticamente. No se utiliza SQL nativo (@Query manual) en ningún repositorio, eliminando el riesgo de inyección SQL.""")

print("Capítulo 5 completado...")

output_path = 'C:\\Users\\sungr\\OneDrive\\Escritorio\\TFG_MockAgent\\memoria\\memoria-tfg.docx'
doc.save(output_path)
print(f"Documento actualizado en: {output_path}")
